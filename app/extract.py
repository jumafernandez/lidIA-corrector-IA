"""Extracción de texto de PDF, DOCX y texto plano."""
import io
import json
import re

MAX_CHARS = 24_000  # ~9k tokens: suficiente para un TFI, evita facturas sorpresa
MAX_UPLOAD_BYTES = 15 * 1024 * 1024
# Un notebook con gráficos pesa mucho más que un informe, y casi todo ese peso son
# imágenes en base64 que descartamos al leerlo. Rechazarlo por tamaño sería rechazar
# justo el entregable típico de un trabajo con datos.
MAX_UPLOAD_NOTEBOOK = 40 * 1024 * 1024


class ExtractionError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> tuple[str, bool]:
    """Devuelve (texto, truncado)."""
    name = (filename or "").lower()
    tope = MAX_UPLOAD_NOTEBOOK if name.endswith(".ipynb") else MAX_UPLOAD_BYTES
    if len(data) > tope:
        raise ExtractionError(f"El archivo supera el máximo de {tope // (1024 * 1024)} MB.")

    if name.endswith(".pdf"):
        text = _from_pdf(data)
    elif name.endswith(".docx"):
        text = _from_docx(data)
    elif name.endswith(".ipynb"):
        text = _from_ipynb(data)
    elif name.endswith((".html", ".htm")):
        text = _from_html(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ExtractionError("Formato no soportado. Subí un PDF, un Word, un notebook (.ipynb), una página HTML o texto plano.")

    text = text.strip()
    if not text:
        raise ExtractionError(
            "No se pudo extraer texto del archivo. Si es un PDF escaneado, exportalo con texto seleccionable."
        )
    if len(text) > MAX_CHARS:
        return text[:MAX_CHARS], True
    return text, False


# Un PDF no guarda párrafos ni encabezados: guarda órdenes de dibujar texto en una
# posición. Todo lo que sigue reconstruye desde esas posiciones lo que la persona ve.
# (Un DOCX sí trae las dos cosas como objetos, y por eso da un resultado mejor.)

BANDA_ALTA = 0.80    # arriba de este porcentaje de la hoja empieza el margen superior
BANDA_BAJA = 0.12    # abajo de este, el inferior
TOLERANCIA_Y = 1.5   # dos fragmentos a menos de esto son el mismo renglón
PROPORCION = 0.6     # en cuántas hojas tiene que repetirse un renglón para ser membrete
BORDE = 3            # cuántos renglones del principio y del final pueden serlo
SALTO_PARRAFO = 1.6  # separación, en veces el interlineado normal, que corta un párrafo


# Cuántas figuras se le mandan al modelo como máximo y a qué tamaño. El tope existe para
# que el costo de una corrección no dependa de cuántos gráficos puso el estudiante; el
# tamaño, porque una captura de 4000 píxeles se paga entera y se ve igual reducida.
TOPE_IMAGENES = 8
LADO_MAX = 1024
MINIMO_UTIL = 120     # por debajo de esto es un ícono o una viñeta, no una figura


def imagenes_de(filename: str, data: bytes) -> list:
    """Las figuras del archivo, listas para mandar: [(mime, bytes)].

    Se extraen las imágenes embebidas y no las páginas rasterizadas: de un informe de 15
    páginas, 12 son texto que el modelo ya recibe como texto, y pagarlas como imagen sería
    pagar dos veces por lo mismo.

    Salen sueltas, sin su epígrafe ni el párrafo que las rodea; van en orden de aparición
    y así se le presentan al modelo. Ante cualquier problema devuelve lista vacía: es una
    mejora de la corrección, no algo sin lo cual no se pueda corregir.
    """
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            crudas = _crudas_pdf(data)
        elif name.endswith(".docx"):
            crudas = _crudas_docx(data)
        elif name.endswith(".ipynb"):
            crudas = _crudas_ipynb(data)
        else:
            return []
    except Exception:  # noqa: BLE001
        return []
    import hashlib

    salida, vistas = [], set()
    for bruto in crudas:
        # El membrete de la institución está en todas las páginas: es una imagen sola
        # repetida, y mandarla ocho veces sería pagar ocho veces por el mismo logo.
        huella = hashlib.sha1(bruto).hexdigest()
        if huella in vistas:
            continue
        vistas.add(huella)
        lista = _achicar(bruto)
        if lista:
            salida.append(lista)
        if len(salida) >= TOPE_IMAGENES:
            break
    return salida


def _achicar(bruto: bytes) -> tuple | None:
    """Reduce la imagen a un tamaño razonable y la devuelve como (mime, bytes)."""
    from PIL import Image

    try:
        im = Image.open(io.BytesIO(bruto))
        im.load()
    except Exception:  # noqa: BLE001
        return None
    if min(im.size) < MINIMO_UTIL:
        return None
    if max(im.size) > LADO_MAX:
        im.thumbnail((LADO_MAX, LADO_MAX))
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=82, optimize=True)
    return "image/jpeg", buf.getvalue()


def _crudas_pdf(data: bytes) -> list:
    from pypdf import PdfReader

    fuera = []
    for pagina in PdfReader(io.BytesIO(data)).pages:
        for im in pagina.images:
            fuera.append(im.data)
            if len(fuera) > TOPE_IMAGENES * 3:   # con margen: después se filtran las chicas
                return fuera
    return fuera


def _crudas_docx(data: bytes) -> list:
    import zipfile

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        nombres = sorted(n for n in zf.namelist() if n.startswith("word/media/"))
        return [zf.read(n) for n in nombres[: TOPE_IMAGENES * 3]]


def _crudas_ipynb(data: bytes) -> list:
    import base64

    nb = json.loads(data.decode("utf-8", errors="replace"))
    fuera = []
    for celda in nb.get("cells") or []:
        for out in celda.get("outputs") or []:
            for k, v in (out.get("data") or {}).items():
                if k.startswith("image/"):
                    try:
                        fuera.append(base64.b64decode(v if isinstance(v, str) else "".join(v)))
                    except Exception:  # noqa: BLE001
                        pass
    return fuera


# Corrección por imagen: se le manda al modelo la página tal como se ve, en vez del texto.
# Así llegan las figuras, las tablas y la maqueta, que en la extracción de texto se pierden.
# El tope existe para que el costo no dependa del largo del trabajo: doce páginas cubren un
# trabajo final completo, y de ahí para arriba conviene la corrección por texto.
TOPE_PAGINAS = 12
ANCHO_PAGINA = 1400   # suficiente para leer cuerpo 11 sin esfuerzo; medido, no estimado


def paginas_de_pdf(data: bytes, tope: int = TOPE_PAGINAS) -> tuple[list, int]:
    """Las páginas del PDF como imágenes. Devuelve ([(mime, bytes)], total_de_paginas).

    Solo PDF: en un Word no existe la página hasta que algo lo maqueta, y maquetarlo
    exigiría LibreOffice en el servidor. Por eso las instancias que corrigen por imagen
    piden la entrega en PDF, en vez de tener un camino distinto por cada formato.
    """
    import pypdfium2 as pdfium

    doc = pdfium.PdfDocument(data)
    total = len(doc)
    paginas = []
    for i in range(min(total, tope)):
        pagina = doc[i]
        escala = ANCHO_PAGINA / max(1.0, pagina.get_width())
        pil = pagina.render(scale=escala).to_pil()
        buf = io.BytesIO()
        pil.convert("RGB").save(buf, format="JPEG", quality=80, optimize=True)
        paginas.append(("image/jpeg", buf.getvalue()))
    return paginas, total


def contar_imagenes(filename: str, data: bytes) -> int:
    """Cuántas imágenes trae el archivo. 0 si no tiene o si no se puede saber.

    La corrección es sobre el texto: un gráfico, un esquema o una captura no llegan al
    modelo. Contarlas no arregla eso, pero permite decirlo —que es distinto de que el
    docente lo descubra cuando la devolución no menciona la matriz de confusión que el
    trabajo mostraba en una figura—.

    Ante cualquier problema devuelve 0: es un aviso, no vale hacer fallar una entrega
    porque el conteo no salió.
    """
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            return _imagenes_pdf(data)
        if name.endswith(".docx"):
            return _imagenes_docx(data)
        if name.endswith(".ipynb"):
            return _imagenes_ipynb(data)
        if name.endswith((".html", ".htm")):
            return _imagenes_html(data)
    except Exception:  # noqa: BLE001
        return 0
    return 0


def _imagenes_pdf(data: bytes) -> int:
    from pypdf import PdfReader

    vistas, total = set(), 0
    for pagina in PdfReader(io.BytesIO(data)).pages:
        recursos = (pagina.get("/Resources") or {}).get("/XObject")
        if not recursos:
            continue
        for nombre, ref in recursos.get_object().items():
            obj = ref.get_object()
            if obj.get("/Subtype") != "/Image":
                continue
            # Un logo repetido en todas las páginas es un objeto solo: se cuenta una vez.
            clave = (getattr(ref, "idnum", None), nombre)
            if clave in vistas:
                continue
            vistas.add(clave)
            total += 1
    return total


def _imagenes_docx(data: bytes) -> int:
    import zipfile

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return sum(1 for n in zf.namelist()
                   if n.startswith("word/media/") and not n.endswith("/"))


def _imagenes_html(data: bytes) -> int:
    import re as _re
    # No existe otra etiqueta HTML que empiece con «img», así que alcanza con buscar eso y
    # se evita una expresión con escapes que es fácil escribir mal.
    return len(_re.findall(rb"<img", data, _re.IGNORECASE))


def _imagenes_ipynb(data: bytes) -> int:
    nb = json.loads(data.decode("utf-8", errors="replace"))
    total = 0
    for celda in nb.get("cells") or []:
        for out in celda.get("outputs") or []:
            datos = out.get("data") or {}
            total += sum(1 for k in datos if k.startswith("image/"))
    return total


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        hojas = [_hoja(p) for p in reader.pages]
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"No se pudo leer el PDF: {exc}") from exc

    if not any(h["renglones"] for h in hojas):
        # Sin posiciones utilizables (PDF raro, o pypdf que no llama al visitante):
        # se vuelve a la extracción lineal y se rearma a ojo.
        try:
            reader = PdfReader(io.BytesIO(data))
            return "\n\n".join(_reflujo(p.extract_text() or "") for p in reader.pages).strip()
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError(f"No se pudo leer el PDF: {exc}") from exc

    _quitar_membrete(hojas)
    bloques = []
    for h in hojas:
        bloques.extend(_parrafos(h["renglones"]))
    return "\n\n".join(b for b in bloques if b)


def _hoja(pagina) -> dict:
    """Renglones de una página, con su altura, reconstruidos desde las coordenadas.

    El visitante recibe cada fragmento con dos matrices: `cm` es la transformación del
    lienzo y `tm` la del texto. La posición real es la composición de ambas; usar `tm`
    sola da coordenadas que no corresponden a la hoja.
    """
    crudos: list = []

    def visitante(texto, cm, tm, fuente, tam):
        if not texto.strip():
            return
        x = cm[0] * tm[4] + cm[2] * tm[5] + cm[4]
        y = cm[1] * tm[4] + cm[3] * tm[5] + cm[5]
        crudos.append((y, x, texto))

    try:
        pagina.extract_text(visitor_text=visitante)
        alto = float(pagina.mediabox.height) or 842.0
    except Exception:  # noqa: BLE001
        return {"alto": 842.0, "renglones": []}

    # Agrupar por altura: los fragmentos de un mismo renglón comparten `y` salvo redondeo.
    grupos: list = []
    for y, x, t in sorted(crudos, key=lambda c: -c[0]):
        if grupos and abs(grupos[-1][0] - y) <= TOLERANCIA_Y:
            grupos[-1][1].append((x, t))
        else:
            grupos.append((y, [(x, t)]))

    renglones = []
    for y, partes in grupos:
        texto = " ".join(t for _, t in sorted(partes)).strip()
        texto = " ".join(texto.split())
        if texto:
            renglones.append({"y": y, "texto": texto})
    return {"alto": alto, "renglones": renglones}


def _quitar_membrete(hojas: list) -> None:
    """Saca encabezado y pie, in situ.

    Se piden tres condiciones a la vez, y las tres son necesarias:

    1. Estar en el margen, arriba o abajo. Un membrete no aparece en el medio.
    2. Ocupar la misma altura en todas las hojas. Un membrete no se mueve.
    3. Decir exactamente lo mismo en todas. Esta es la que evita el desastre: si un
       programa tiene margen superior chico, el cuerpo empieza dentro de la banda y
       siempre a la misma altura, así que las dos primeras condiciones se cumplirían
       y se borraría la primera línea de cada página. Como el cuerpo dice algo
       distinto en cada hoja, exigir el texto idéntico lo deja intacto.

    El costo de la tercera es que un encabezado que cambia de contenido página a página
    —el nombre de la unidad, un número de página— no se detecta y queda en el texto.
    Se prefiere ese error: dejar de más es un renglón molesto, borrar de más es perder
    contenido del programa sin que nadie se entere.
    """
    if len(hojas) < 3:
        return
    veces: dict = {}
    for n, h in enumerate(hojas):
        con_texto = [r for r in h["renglones"]]
        borde = con_texto[:BORDE] + con_texto[-BORDE:]
        for r in borde:
            rel = r["y"] / h["alto"]
            if rel >= BANDA_ALTA or rel <= BANDA_BAJA:
                veces.setdefault((round(rel, 2), r["texto"]), set()).add(n)
    minimo = max(2, int(len(hojas) * PROPORCION))
    membrete = {k for k, v in veces.items() if len(v) >= minimo}
    if not membrete:
        return
    for h in hojas:
        h["renglones"] = [r for r in h["renglones"]
                          if (round(r["y"] / h["alto"], 2), r["texto"]) not in membrete]


VINETA = re.compile(r"^([-–—•*·]|\d{1,2}[.)]|[a-zA-Z][.)])\s+")


def _parrafos(renglones: list) -> list:
    """Une los renglones de un mismo párrafo y devuelve la lista de párrafos.

    El corte de renglón dentro de un párrafo es de la caja de texto, no del contenido.
    Los cortes que sí significan algo son tres: un salto vertical mayor que el
    interlineado habitual, una viñeta, y un título en mayúsculas.
    """
    if not renglones:
        return []
    saltos = sorted(abs(renglones[i]["y"] - renglones[i + 1]["y"])
                    for i in range(len(renglones) - 1))
    normal = saltos[len(saltos) // 2] if saltos else 0     # interlineado típico

    parrafos: list = []
    actual: list = []
    for i, r in enumerate(renglones):
        titulo = r["texto"].isupper() and len(r["texto"]) < 90
        lejos = (i > 0 and normal
                 and abs(renglones[i - 1]["y"] - r["y"]) > normal * SALTO_PARRAFO)
        if actual and (lejos or titulo or VINETA.match(r["texto"])
                       or (actual[-1].isupper() and len(actual[-1]) < 90)):
            parrafos.append(" ".join(actual))
            actual = []
        if actual and actual[-1].endswith("-"):
            actual[-1] = actual[-1][:-1] + r["texto"]      # palabra cortada por guión
        else:
            actual.append(r["texto"])
    if actual:
        parrafos.append(" ".join(actual))
    return parrafos


def _reflujo(pagina: str) -> str:
    """Rearma el párrafo cuando la extracción lineal devuelve una palabra por renglón."""
    pagina = pagina.replace("\r\n", "\n").replace("\r", "\n")
    con_texto = [r.strip() for r in pagina.split("\n") if r.strip()]
    if not con_texto:
        return ""
    sueltas = sum(1 for r in con_texto if " " not in r)
    if sueltas / len(con_texto) < 0.7:
        return pagina
    partes: list = []
    for r in con_texto:
        if partes and partes[-1].endswith("-"):
            partes[-1] = partes[-1][:-1] + r
        else:
            partes.append(r)
    return " ".join(partes)


# Un notebook puede traer salidas enormes —una imagen en base64, un DataFrame entero, un
# traceback de veinte líneas—. Se recortan por celda para que el trabajo entre en el
# límite general sin que una sola salida se coma todo el texto del informe.
SALIDA_MAX = 1200
SALIDAS_POR_CELDA = 3


def _from_ipynb(data: bytes) -> str:
    """Texto de un notebook: su prosa, su código y lo que dio al correr.

    Es el entregable habitual de un trabajo con datos, y las tres partes hacen falta para
    corregirlo: en el markdown está el razonamiento, en el código lo que realmente se hizo,
    y en las salidas si eso funcionó y qué números dio. Un notebook sin sus salidas se
    puede leer, pero no se puede saber si lo que afirma es cierto.

    Las imágenes se anotan pero no se transcriben: un gráfico en base64 son cientos de
    miles de caracteres que el modelo no puede ver igual.
    """
    try:
        nb = json.loads(data.decode("utf-8", errors="replace"))
        celdas = nb["cells"]
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"No se pudo leer el notebook: {exc}") from exc

    partes: list = []
    n_codigo = 0
    for celda in celdas:
        tipo = celda.get("cell_type")
        fuente = "".join(celda.get("source") or []).strip()
        if tipo == "markdown":
            if fuente:
                partes.append(fuente)
        elif tipo == "code":
            n_codigo += 1
            if fuente:
                partes.append(f"[Celda {n_codigo}]\n```python\n{fuente}\n```")
            salidas = _salidas(celda.get("outputs") or [])
            if salidas:
                partes.append(f"[Salida de la celda {n_codigo}]\n{salidas}")
    if not partes:
        raise ExtractionError("El notebook no tiene contenido: todas sus celdas están vacías.")
    return "\n\n".join(partes)


def _salidas(outputs: list) -> str:
    """Lo que imprimió cada celda, recortado y sin imágenes."""
    trozos: list = []
    for out in outputs[:SALIDAS_POR_CELDA]:
        clase = out.get("output_type")
        texto = ""
        if clase == "stream":
            texto = "".join(out.get("text") or [])
        elif clase in ("execute_result", "display_data"):
            datos = out.get("data") or {}
            if "text/plain" in datos:
                texto = "".join(datos["text/plain"])
            elif any(k.startswith("image/") for k in datos):
                texto = "(gráfico)"
        elif clase == "error":
            # El error importa: una celda que falla dice más que una que no se ejecutó.
            texto = f"{out.get('ename', 'Error')}: {out.get('evalue', '')}"
        texto = texto.strip()
        if not texto:
            continue
        if len(texto) > SALIDA_MAX:
            texto = texto[:SALIDA_MAX] + "\n… (salida recortada)"
        trozos.append(texto)
    if len(outputs) > SALIDAS_POR_CELDA:
        trozos.append(f"… ({len(outputs) - SALIDAS_POR_CELDA} salidas más)")
    return "\n".join(trozos)


# Etiquetas cuyo contenido no es del documento: el guion que lo hace andar y su estilo.
# Un notebook exportado a HTML trae bastante de las dos cosas.
# Etiquetas cuyo CONTENIDO se descarta. Solo van las que abren y cierran: <meta> y <link>
# no cierran nunca, así que sumaban al contador y no lo bajaban jamás, y a partir de ahí se
# descartaba el documento entero. Tampoco hacían falta: no tienen texto adentro, y las que
# viven en el encabezado ya quedan tapadas por «head».
IGNORADAS = {"script", "style", "head", "noscript"}
# Etiquetas que separan bloques: sin esto el texto sale todo pegado en un renglón.
BLOQUES = {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6",
           "pre", "blockquote", "section", "article", "td", "th"}


def _from_html(data: bytes) -> str:
    """Texto de una página HTML. Es el formato con el que se exporta un notebook.

    Se usa el analizador de la biblioteca estándar en vez de sumar una dependencia: para
    sacar el texto de un documento alcanza, y lo que no alcanzaría —interpretar CSS, correr
    guiones— tampoco haría falta acá.
    """
    from html.parser import HTMLParser

    class Lector(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.partes, self.omitir = [], 0

        def handle_starttag(self, tag, attrs):
            if tag in IGNORADAS:
                self.omitir += 1
            elif tag in BLOQUES:
                self.partes.append("\n")

        def handle_endtag(self, tag):
            if tag in IGNORADAS and self.omitir:
                self.omitir -= 1
            elif tag in BLOQUES:
                self.partes.append("\n")

        def handle_data(self, texto):
            if not self.omitir and texto.strip():
                self.partes.append(texto)

    lector = Lector()
    lector.feed(data.decode("utf-8", errors="replace"))
    crudo = "".join(lector.partes)
    # Renglones sin espacios de sobra, y como mucho un renglón vacío entre bloques.
    lineas, salida = [l.strip() for l in crudo.split("\n")], []
    for l in lineas:
        if l or (salida and salida[-1]):
            salida.append(" ".join(l.split()))
    return "\n".join(salida).strip()


def _from_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"No se pudo leer el DOCX: {exc}") from exc
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
