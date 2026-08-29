"""La plantilla de una instancia de evaluación: cómo se entrega y cómo se lee.

Un docente tiene un documento, no cuatro. Pedirle que lo corte en pedazos para llenar
campo por campo es trasladarle nuestro modelo de datos. Acá se hace al revés: se le
entrega un formulario, lo completa, lo sube entero y el sistema lo separa.

El formulario es un formulario de verdad, no una sugerencia. Las secciones se reconocen
por su título exacto y son obligatorias; si el documento no las trae, se rechaza con el
motivo. No se intenta deducir dónde termina la consigna y empieza la rúbrica: equivocarse
en eso significa corregir contra el criterio equivocado, y eso no se puede permitir a
cambio de ahorrarle una descarga a nadie.
"""
import io
import re
import unicodedata


class FormatoInvalido(Exception):
    """El documento no respeta la plantilla. El mensaje es para el docente."""


# ------------------------------------------------------------------ el contrato

# Qué secciones tiene cada tipo y en qué campo del sistema cae cada una. El orden es el
# del documento. `clave` es cómo la nombramos internamente; el título es lo que se
# escribe en el documento y lo que se busca al leerlo.
SECCIONES = {
    "abierto": [
        ("consigna", "CONSIGNA", "Lo que lee el estudiantado: qué tiene que entregar, en qué\n"
                                 "formato y con qué condiciones."),
        ("rubrica", "RÚBRICA", "Un criterio por línea. Es contra esto que Lidia corrige, y el\n"
                               "estudiantado no lo ve hasta la devolución."),
    ],
    "escrito": [
        ("consigna", "CONSIGNA", "Lo que lee el estudiantado antes de las preguntas: condiciones,\n"
                                 "tiempo, materiales permitidos."),
        ("preguntas", "PREGUNTAS", "Una por línea, numeradas. El puntaje va opcional entre\n"
                                   "corchetes al final; si no está, vale 1.\n"
                                   "Ejemplo:  1. ¿Qué distingue a un agente racional? [2]"),
        ("respuestas", "RESPUESTAS ESPERADAS", "Una por pregunta, con el mismo número. Material\n"
                                               "interno: el estudiantado no lo ve."),
    ],
    "choice": [
        ("consigna", "CONSIGNA", "Lo que lee el estudiantado antes de las preguntas."),
        ("preguntas", "PREGUNTAS", "Numeradas, con sus opciones debajo como a) b) c). El puntaje\n"
                                   "va opcional entre corchetes al final del enunciado.\n"
                                   "Ejemplo:  1. ¿Cuál de estas afirmaciones es correcta? [2]\n"
                                   "            a) Primera opción\n"
                                   "            b) Segunda opción"),
        ("clave", "CLAVE DE CORRECCIÓN", "Una línea por pregunta, con su letra.\n"
                                         "Ejemplo:  1. b"),
    ],
}

NOMBRE_TIPO = {"abierto": "trabajo abierto", "escrito": "examen escrito",
               "choice": "multiple choice"}


MARCA = "==="   # decoración del título, para que se lea como marca y no como texto


def rotulo(titulo: str) -> str:
    """El título tal como se escribe en el documento."""
    return f"{MARCA} {titulo} {MARCA}"


def _plano(texto: str) -> str:
    """Sin tildes, sin puntuación de borde y en mayúsculas, para comparar títulos.

    Un título escrito «Rúbrica», «RUBRICA:» o «— RÚBRICA —» es el mismo título. Lo que no
    se acepta es que no esté, o que diga otra cosa.
    """
    t = unicodedata.normalize("NFD", texto)
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    return t.strip(" \t.:-–—=_*#[]()").upper()


def _titulos(tipo: str) -> dict:
    return {_plano(titulo): clave for clave, titulo, _ in SECCIONES[tipo]}


# ------------------------------------------------------------------ leer el documento

def separar(tipo: str, texto: str) -> dict:
    """Parte el documento en sus secciones. Falla si falta alguna.

    Devuelve {clave: texto}. No interpreta el contenido de cada sección: eso lo hacen
    `preguntas()` y `respuestas()`, que sí conocen el formato de cada una.
    """
    if tipo not in SECCIONES:
        raise FormatoInvalido(f"Tipo de instancia desconocido: {tipo}.")
    titulos = _titulos(tipo)

    encontradas: dict = {}
    actual = None
    for linea in (texto or "").splitlines():
        clave = titulos.get(_plano(linea)) if linea.strip() else None
        if clave:
            actual = clave
            encontradas[actual] = []
            continue
        if actual:
            encontradas[actual].append(linea)

    faltan = [titulo for clave, titulo, _ in SECCIONES[tipo] if clave not in encontradas]
    if faltan:
        cuales = ", ".join(f"«{t}»" for t in faltan)
        raise FormatoInvalido(
            f"El documento no respeta la plantilla de {NOMBRE_TIPO[tipo]}: "
            f"no encontré {'la sección' if len(faltan) == 1 else 'las secciones'} {cuales}. "
            "Descargá la plantilla y completala sin borrar ni renombrar los títulos."
        )

    partes = {c: "\n".join(v).strip() for c, v in encontradas.items()}
    vacias = [t for c, t, _ in SECCIONES[tipo] if not partes.get(c)]
    if vacias:
        cuales = ", ".join(f"«{t}»" for t in vacias)
        raise FormatoInvalido(
            f"{'La sección' if len(vacias) == 1 else 'Las secciones'} {cuales} "
            f"{'quedó vacía' if len(vacias) == 1 else 'quedaron vacías'} en el documento."
        )
    return partes


# Una pregunta empieza con su número: «1.», «1)» o «Pregunta 1:». El puntaje, si está,
# va entre corchetes al final del enunciado.
NUMERADA = re.compile(r"^\s*(?:pregunta\s*)?(\d{1,3})\s*[.)\-:]\s*(.+)$", re.IGNORECASE)
OPCION = re.compile(r"^\s*([a-hA-H])\s*[.)\-]\s*(.+)$")
PUNTAJE = re.compile(r"\s*\[\s*(\d+(?:[.,]\d+)?)\s*(?:puntos?|pts?\.?)?\s*\]\s*$", re.IGNORECASE)


def preguntas(texto: str, con_opciones: bool) -> list:
    """Las preguntas de la sección PREGUNTAS, en orden.

    Devuelve [{orden, enunciado, opciones, puntaje}]. `opciones` es el texto con una
    opción por línea, tal como lo espera el formulario.
    """
    items: list = []
    siguiente = 1
    for linea in texto.splitlines():
        if not linea.strip():
            continue
        m = NUMERADA.match(linea)
        # Igual que en las respuestas: una pregunta arranca en el número que corresponde,
        # así que una enumeración dentro de un enunciado no la parte en dos.
        if m and int(m.group(1)) != siguiente:
            m = None
        if m:
            siguiente += 1
            enunciado = m.group(2).strip()
            puntaje = 1.0
            p = PUNTAJE.search(enunciado)
            if p:
                puntaje = float(p.group(1).replace(",", "."))
                enunciado = enunciado[:p.start()].strip()
            items.append({"orden": int(m.group(1)), "enunciado": enunciado,
                          "opciones": [], "puntaje": puntaje})
            continue
        o = OPCION.match(linea)
        if o and con_opciones and items:
            items[-1]["opciones"].append(o.group(2).strip())
            continue
        if items:                      # continuación del enunciado anterior
            items[-1]["enunciado"] = (items[-1]["enunciado"] + " " + linea.strip()).strip()
        else:
            raise FormatoInvalido(
                "La sección «PREGUNTAS» tiene texto antes de la primera pregunta numerada. "
                "Cada pregunta tiene que empezar con su número (por ejemplo «1.»)."
            )

    if not items:
        raise FormatoInvalido("No encontré ninguna pregunta numerada en la sección «PREGUNTAS».")
    _sin_saltos(items)
    for it in items:
        it["opciones"] = "\n".join(it["opciones"])
    return items


def respuestas(texto: str, esperadas: int, como_letra: bool) -> dict:
    """Las respuestas por número de pregunta. Falla si falta alguna o sobra.

    `como_letra` es el multiple choice: ahí la respuesta es la letra de la opción correcta
    y cualquier otra cosa es un error de carga, no una respuesta libre.
    """
    # Una respuesta empieza donde aparece el número que sigue, y no en cualquier renglón
    # numerado: una respuesta bien escrita puede tener adentro su propia lista —«1. …,
    # 2. …»— y esos renglones son parte de la respuesta, no el comienzo de otra.
    por_numero: dict = {}
    ultimo = None
    siguiente = 1
    for linea in texto.splitlines():
        if not linea.strip():
            if ultimo is not None:
                por_numero[ultimo] += "\n"
            continue
        m = NUMERADA.match(linea)
        if m and int(m.group(1)) == siguiente:
            ultimo = siguiente
            siguiente += 1
            por_numero[ultimo] = m.group(2).strip()
        elif ultimo is not None:
            por_numero[ultimo] = (por_numero[ultimo].rstrip() + "\n" + linea.strip()
                                  if por_numero[ultimo].endswith("\n")
                                  else por_numero[ultimo] + " " + linea.strip()).strip()
        else:
            raise FormatoInvalido(
                "Las respuestas tienen que ir numeradas igual que las preguntas "
                "(por ejemplo «1.»); encontré texto antes de la primera."
            )

    faltan = [n for n in range(1, esperadas + 1) if n not in por_numero]
    if faltan:
        cuales = ", ".join(str(n) for n in faltan)
        raise FormatoInvalido(
            f"Falta la respuesta de {'la pregunta' if len(faltan) == 1 else 'las preguntas'} {cuales}."
        )
    sobran = [n for n in por_numero if n > esperadas]
    if sobran:
        cuales = ", ".join(str(n) for n in sorted(sobran))
        raise FormatoInvalido(
            f"Hay respuestas para {cuales}, pero no hay tantas preguntas."
        )

    if como_letra:
        malas = [str(n) for n, v in por_numero.items() if not re.fullmatch(r"[a-hA-H]", v.strip(" .)"))]
        if malas:
            raise FormatoInvalido(
                f"En la clave, {'la respuesta' if len(malas) == 1 else 'las respuestas'} "
                f"{', '.join(malas)} no {'es' if len(malas) == 1 else 'son'} una letra de opción."
            )
        por_numero = {n: v.strip(" .)").lower() for n, v in por_numero.items()}
    return por_numero


def _sin_saltos(items: list) -> None:
    """Las preguntas tienen que ser 1, 2, 3… Un salto casi siempre es una que se perdió."""
    numeros = [i["orden"] for i in items]
    esperado = list(range(1, len(numeros) + 1))
    if numeros != esperado:
        raise FormatoInvalido(
            f"Las preguntas están numeradas {', '.join(map(str, numeros))} y tendrían que ir "
            f"de 1 a {len(numeros)} sin saltos ni repetidos."
        )


def leer(tipo: str, texto: str) -> dict:
    """Todo junto: del documento a los campos del sistema.

    Devuelve {"consigna", "rubrica", "items"}. `items` es [] en el trabajo abierto.
    """
    partes = separar(tipo, texto)
    if tipo == "abierto":
        return {"consigna": partes["consigna"], "rubrica": partes["rubrica"], "items": []}

    choice = tipo == "choice"
    items = preguntas(partes["preguntas"], con_opciones=choice)
    respu = respuestas(partes["clave" if choice else "respuestas"], len(items), como_letra=choice)
    for it in items:
        it["respuesta"] = respu[it["orden"]]
    if choice:
        sin_opciones = [str(i["orden"]) for i in items if not i["opciones"].strip()]
        if sin_opciones:
            raise FormatoInvalido(
                f"{'La pregunta' if len(sin_opciones) == 1 else 'Las preguntas'} "
                f"{', '.join(sin_opciones)} no {'tiene' if len(sin_opciones) == 1 else 'tienen'} "
                "opciones debajo (a, b, c…)."
            )
    return {"consigna": partes["consigna"], "rubrica": "", "items": items}


# ------------------------------------------------------------------ entregar la plantilla

def justificar(parrafo) -> None:
    """Alinea el párrafo a los dos márgenes.

    Es solo estética del documento que entregamos. En Word no afecta la lectura: el texto
    se saca de los párrafos, no de cómo están dibujados. (En un PDF sí importaría, porque
    ahí la justificación mete espacios de relleno que hay que limpiar al extraer.)
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _resaltar(parrafo) -> None:
    """Le pone fondo y borde al título, para que se vea como una marca del formulario.

    Es solo apariencia: al extraer el texto no queda rastro de esto, así que no cambia en
    nada cómo se lee el documento. Existe para la persona, no para el sistema.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    pPr = parrafo._p.get_or_add_pPr()
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), "E8F2F3")
    pPr.append(sombra)
    bordes = OxmlElement("w:pBdr")
    for lado in ("top", "bottom"):
        b = OxmlElement(f"w:{lado}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), "0F6E78")
        bordes.append(b)
    pPr.append(bordes)
    for run in parrafo.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x78)


def plantilla_docx(tipo: str) -> bytes:
    """El formulario en blanco, en Word.

    Se genera desde el mismo `SECCIONES` que después se lee, así que el documento que
    entregamos y el que sabemos leer no pueden separarse con el tiempo.

    Toda la ayuda va ARRIBA del primer título, no debajo de cada uno. Al leer, lo que está
    antes del primer título se descarta, así que la plantilla sin completar se rechaza
    sola: si la ayuda estuviera dentro de las secciones, subirla en blanco cargaría el
    instructivo como si fuera la consigna.
    """
    from docx import Document
    from docx.shared import Pt

    if tipo not in SECCIONES:
        raise FormatoInvalido(f"Tipo de instancia desconocido: {tipo}.")

    doc = Document()
    doc.add_heading(f"Plantilla de {NOMBRE_TIPO[tipo]} · LidIA", level=0)
    intro = doc.add_paragraph(
        "Completá cada sección debajo de su título y borrá estas instrucciones si querés. "
        "Los títulos en mayúsculas son los que el sistema busca para separar el documento: "
        "no los borres ni los renombres. Tipografía, colores y márgenes son libres."
    )
    intro.runs[0].font.size = Pt(10)
    justificar(intro)
    for _, titulo, ayuda in SECCIONES[tipo]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{titulo}: ").bold = True
        r = p.add_run(" ".join(ayuda.split()))
        r.font.size = Pt(10)
        justificar(p)

    for _, titulo, _ in SECCIONES[tipo]:
        doc.add_paragraph()
        h = doc.add_heading(rotulo(titulo), level=1)
        _resaltar(h)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def plantilla_txt(tipo: str) -> str:
    """La misma plantilla en texto plano, para quien no usa Word."""
    if tipo not in SECCIONES:
        raise FormatoInvalido(f"Tipo de instancia desconocido: {tipo}.")
    lineas = [f"PLANTILLA DE {NOMBRE_TIPO[tipo].upper()} · LidIA",
              "=" * (len(NOMBRE_TIPO[tipo]) + 22), "",
              "Completá cada sección debajo de su título. Los títulos en mayúsculas son los",
              "que el sistema busca para separar el documento: no los borres ni los renombres.",
              ""]
    for _, titulo, ayuda in SECCIONES[tipo]:
        lineas.append(f"- {titulo}: " + " ".join(ayuda.split()))
    for _, titulo, _ in SECCIONES[tipo]:
        lineas += ["", "", rotulo(titulo), ""]
    return "\n".join(lineas) + "\n"
